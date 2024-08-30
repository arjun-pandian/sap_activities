from pywinauto import application, timings
import pyautogui
import time
from docx import Document
import subprocess
import pandas as pd
import win32com.client
from datetime import datetime 
from docx.shared import Inches
from datetime import date
import os
import sys

current_directory = os.getcwd()
functions_directory = os.path.abspath(os.path.join(current_directory, 'Files', 'Common'))
sys.path.append(functions_directory)

from sap_functions import sap_login, create_session, take_and_save_screenshot
import time
import pyautogui
import re


#P08-100 System
sap_login("P08", "100")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsick"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "p08_1_sick")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm58"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/txtBENUTZER-LOW").text = ""
session.findById("wnd[0]/usr/txtBENUTZER-LOW").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_2_sm58")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm66"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "p08_3_sm66")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm51"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "p08_4_sm51")
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").selectedRows = "0"
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").doubleClickCurrentCell()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,224)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "p08_5_sm51")
session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").currentCellRow = 1
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").selectedRows = "1"
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").doubleClickCurrentCell()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,246)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "p08_6_sm51")
session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").currentCellRow = 2
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").selectedRows = "2"
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[1]/shell").doubleClickCurrentCell()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,234)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "p08_7_sm51")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm50"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,237)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "p08_8_sm50")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nst22"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "p08_9_st22")

session.findById("wnd[0]/usr/txtS_UNAME-LOW").text = "*"
session.findById("wnd[0]/usr/txtS_UNAME-LOW").setFocus()
session.findById("wnd[0]/usr/txtS_UNAME-LOW").caretPosition = 1
session.findById("wnd[0]/usr/btnTODAY").press()
take_and_save_screenshot("GSA", "p08_10_st22")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/btnYESTERD").press()
take_and_save_screenshot("GSA", "p08_11_st22")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm21"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_12_sm21")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").text = "p08db_P08_01"
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").setFocus()
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").caretPosition = 12
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_13_sm21")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").text = "frasap28_P08_02"
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").setFocus()
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").caretPosition = 15
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_14_sm21")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").text = "frasap29_P08_02"
session.findById("wnd[0]/usr/ctxtINSTOPT-LOW").setFocus()
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_15_sm21")

session.findById("wnd[0]/tbar[0]/okcd").text = "/ndb12"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/btnTOTAL_BUTTON").press()
take_and_save_screenshot("GSA", "p08_16_db12")

session.findById("wnd[0]/tbar[0]/okcd").text = "/ndb01"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "p08_17_db01")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm12"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/txtSEQG3-GUNAME").text = "*"
session.findById("wnd[0]/usr/txtSEQG3-GUNAME").setFocus()
session.findById("wnd[0]/usr/txtSEQG3-GUNAME").caretPosition = 1
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_18_sm12")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm13"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/radSEL_STOPPED").select()
session.findById("wnd[0]/usr/txtFROM_DATE").text = ""
session.findById("wnd[0]/usr/radSEL_STOPPED").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/txtFROM_DATE").setFocus()
session.findById("wnd[0]/usr/txtFROM_DATE").caretPosition = 0
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_19_sm13")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm37"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/chkBTCH2170-SCHEDUL").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-READY").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-FINISHED").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-ABORTED").selected = False
session.findById("wnd[0]/usr/txtBTCH2170-USERNAME").text = "*"
session.findById("wnd[0]/usr/txtBTCH2170-USERNAME").setFocus()
session.findById("wnd[0]/usr/txtBTCH2170-USERNAME").caretPosition = 1
session.findById("wnd[0]/tbar[1]/btn[8]").press()
session.findById("wnd[0]/usr").verticalScrollbar.position = 1
session.findById("wnd[0]/usr").verticalScrollbar.position = 2
take_and_save_screenshot("GSA", "p08_20_sm37")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/chkBTCH2170-ABORTED").selected = True
session.findById("wnd[0]/usr/chkBTCH2170-RUNNING").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-RUNNING").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "p08_21_sm37")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsost"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/subSUB:SAPLSBCS_OUT:1100/subTOPSUB:SAPLSBCS_OUT:1110/tabsTAB1/tabpTAB1_FC1/ssubTAB1_SCA:SAPLSBCS_OUT:0003/btnREFRICO2").press()
session.findById("wnd[0]/usr/subSUB:SAPLSBCS_OUT:1100/subMAINSUB:SAPLSBCS_OUT:1120/cntlCUSTOM_CONTROL/shellcont/shell/shellcont[0]/shell").setCurrentCell (-1,"STAT_DATE")
session.findById("wnd[0]/usr/subSUB:SAPLSBCS_OUT:1100/subMAINSUB:SAPLSBCS_OUT:1120/cntlCUSTOM_CONTROL/shellcont/shell/shellcont[0]/shell").selectColumn ("STAT_DATE")
session.findById("wnd[0]/usr/subSUB:SAPLSBCS_OUT:1100/subMAINSUB:SAPLSBCS_OUT:1120/cntlCUSTOM_CONTROL/shellcont/shell/shellcont[0]/shell").contextMenu()
session.findById("wnd[0]/usr/subSUB:SAPLSBCS_OUT:1100/subMAINSUB:SAPLSBCS_OUT:1120/cntlCUSTOM_CONTROL/shellcont/shell/shellcont[0]/shell").selectContextMenuItem ("&SORT_ASC")
take_and_save_screenshot("GSA", "p08_22_sost")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

#PKQ-100 System
sap_login("PKQ", "100")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsick"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_1_sick")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm58"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/txtBENUTZER-LOW").text = "*"
session.findById("wnd[0]/usr/txtBENUTZER-LOW").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_2_sm58")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm66"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_3_sm66")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm51"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_4_sm51")

session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[0]/shell").selectedRows = "0"
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell/shellcont[0]/shell").doubleClickCurrentCell()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,275)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "pkq_5_sm51")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm50"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,266)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("GSA", "pkq_6_sm50")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nst22"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_7_st22")

session.findById("wnd[0]/usr/btnTODAY").press()
take_and_save_screenshot("GSA", "pkq_8_st22")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/txtS_UNAME-LOW").text = "*"
session.findById("wnd[0]/usr/txtS_UNAME-LOW").setFocus()
session.findById("wnd[0]/usr/txtS_UNAME-LOW").caretPosition = 1
session.findById("wnd[0]/usr/btnYESTERD").press()
take_and_save_screenshot("GSA", "pkq_9_st22")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm21"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_10_sm21")

session.findById("wnd[0]/tbar[0]/okcd").text = "/ndb12"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/btnTOTAL_BUTTON").press()
take_and_save_screenshot("GSA", "pkq_11_db12")

session.findById("wnd[0]/tbar[0]/okcd").text = "/ndb01"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_12_db01")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm12"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/txtSEQG3-GUNAME").text = "*"
session.findById("wnd[0]/usr/txtSEQG3-GUNAME").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_13_sm12")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm13"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/radSEL_STOPPED").select()
session.findById("wnd[0]/usr/radSEL_STOPPED").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_14_sm13")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsm37"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/chkBTCH2170-SCHEDUL").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-READY").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-FINISHED").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-ABORTED").selected = False
session.findById("wnd[0]/usr/txtBTCH2170-USERNAME").text = "*"
session.findById("wnd[0]/usr/chkBTCH2170-ABORTED").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_15_sm37")

session.findById("wnd[0]/tbar[0]/btn[3]").press()
session.findById("wnd[0]/usr/chkBTCH2170-ABORTED").selected = True
session.findById("wnd[0]/usr/chkBTCH2170-RUNNING").selected = False
session.findById("wnd[0]/usr/chkBTCH2170-RUNNING").setFocus()
session.findById("wnd[0]/tbar[1]/btn[8]").press()
take_and_save_screenshot("GSA", "pkq_16_sm37")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nsost"
session.findById("wnd[0]").sendVKey (0)
take_and_save_screenshot("GSA", "pkq_17_sost")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

base_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Screenshots'))

gsa_screenshot_names = [
    "p08_1_sick.png", "p08_2_sm58.png", "p08_3_sm66.png", "p08_4_sm51.png", 
    "p08_5_sm51.png", "p08_6_sm51.png", "p08_7_sm51.png", "p08_8_sm50.png",
    "p08_9_st22.png", "p08_10_st22.png", "p08_11_st22.png", "p08_12_sm21.png",
    "p08_13_sm21.png", "p08_14_sm21.png", "p08_15_sm21.png", "p08_16_db12.png",
    "p08_17_db01.png", "p08_18_sm12.png", "p08_19_sm13.png", "p08_20_sm37.png",
    "p08_21_sm37.png", "p08_22_sost.png", "pkq_1_sick.png", "pkq_2_sm58.png",
    "pkq_3_sm66.png", "pkq_4_sm51.png", "pkq_5_sm51.png", "pkq_6_sm50.png",
    "pkq_7_st22.png", "pkq_8_st22.png", "pkq_9_st22.png", "pkq_10_sm21.png",
    "pkq_11_db12.png", "pkq_12_db01.png", "pkq_13_sm12.png", "pkq_14_sm13.png",
    "pkq_15_sm37.png", "pkq_16_sm37.png", "pkq_17_sost.png"
]

all_screenshot_paths = [
    f"{base_path}/GSA/{screenshot}" for screenshot in gsa_screenshot_names
]

today = date.today().strftime("%d.%m.%Y")
template_path = os.path.abspath(os.path.join(current_directory, 'Files','GSA','GSA_Monitoring_Template.docx'))
documents_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Documents'))
if not os.path.exists(documents_path):
    os.makedirs(documents_path)

output_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Documents',f'GSA_Monitoring_{today}.docx'))
doc = Document(template_path)

for i, screenshot_file in enumerate(all_screenshot_paths):
    placeholder = f"#IMAGE{i+1}#"
    for para in doc.paragraphs:
            if placeholder in para.text:
                index = para.text.find(placeholder)
                para.text = para.text.replace(placeholder,"")
                run = para.add_run()
                run.add_picture(screenshot_file, width=Inches(6))
                run.add_break()
                break

today = date.today().strftime("(%d/%m/%Y)")
for para in doc.paragraphs:
    if "#DATE#" in para.text:
        for run in para.runs:
            if "#DATE#" in run.text:
                run.text = run.text.replace("#DATE#", "")
                run.bold = True
                break
        run = para.add_run(today)
        run.bold = True


doc.save(output_path)

print(f"Word document with screenshots saved at {output_path}")