from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import date
import sys
import os
current_directory = os.getcwd()


base_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Screenshots'))


pk9_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

pr3_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

pmp_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

pks_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_al08.png", "14_al11.png"
]

pkx_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_al08.png", "14_al11.png"
]

pl5_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png",
    "8_smq1.png", "9_stms_1.png", "10_stms_2.png", "11_sm58.png",
    "12_db02.png", "13_db12.png", "14_al08.png", "15_al11.png",
    "16_tablespace.png"
]

cp5_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

pm6_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

pr5_screenshot_names = [
    "1_sm12.png", "2_sm13.png", "3_sm21.png", "4_st22.png",
    "5_sm37_1.png", "6_sm37_2.png", "7_sm51.png", "8_sm66.png",
    "9_smq1.png", "10_stms_1.png", "11_stms_2.png", "12_sm58.png",
    "13_db02.png", "14_db12.png", "15_al08.png", "16_al11.png",
    "17_tablespace.png"
]

all_screenshot_paths = [
    f"{base_path}/pr3/{screenshot}" for screenshot in pr3_screenshot_names
] + [
    f"{base_path}/pk9/{screenshot}" for screenshot in pk9_screenshot_names
] + [
    f"{base_path}/pmp/{screenshot}" for screenshot in pmp_screenshot_names
] + [
    f"{base_path}/pkx/{screenshot}" for screenshot in pkx_screenshot_names
] + [
    f"{base_path}/pl5/{screenshot}" for screenshot in pl5_screenshot_names
] + [
    f"{base_path}/cp5/{screenshot}" for screenshot in cp5_screenshot_names
] + [
    f"{base_path}/pm6/{screenshot}" for screenshot in pm6_screenshot_names
] + [
    f"{base_path}/pks/{screenshot}" for screenshot in pks_screenshot_names
] + [
    f"{base_path}/pr5/{screenshot}" for screenshot in pr5_screenshot_names
]

def generate_document():
    today = date.today().strftime("%d.%m.%Y")
    template_path = os.path.abspath(os.path.join(current_directory, 'Files', 'Manual_Monitoring','S2_Manual_Monitoring_template.docx'))
    documents_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Documents'))
    if not os.path.exists(documents_path):
        os.makedirs(documents_path)

    output_path = os.path.abspath(os.path.join(documents_path, f'S2_Manual_Monitoring_{today}.docx'))
    

    doc = Document(template_path)

    for i, screenshot_file in enumerate(all_screenshot_paths):
        placeholder = f"#IMAGE{i+1}#"
        for para in doc.paragraphs:
                if placeholder in para.text:
                    index = para.text.find(placeholder)
                    para.text = para.text.replace(placeholder,"")
                    run = para.add_run()
                    run.add_picture(screenshot_file, width=Inches(6))
                    run.add_break()
                    break
    
    today = date.today().strftime("(%d/%m/%Y)")
    for para in doc.paragraphs:
        if "#DATE#" in para.text:
            for run in para.runs:
                if "#DATE#" in run.text:
                    run.text = run.text.replace("#DATE#", "")
                    run.bold = True
                    break
            run = para.add_run(today)
            run.bold = True


    doc.save(output_path)

    print(f"Word document with screenshots saved at {output_path}")


if __name__ == "__main__":

    generate_document()