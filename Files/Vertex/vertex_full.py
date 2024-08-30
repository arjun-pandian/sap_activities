from pywinauto import application, timings
import pyautogui
import time
from docx import Document
import subprocess
import pandas as pd
import win32com.client
from datetime import datetime 
from docx.shared import Inches
from datetime import date
import os
import sys

current_directory = os.getcwd()
functions_directory = os.path.abspath(os.path.join(current_directory, 'Files', 'Common'))
sys.path.append(functions_directory)

from sap_functions import sap_login, create_session, take_and_save_screenshot
import time
import pyautogui
import re


#VA5-700 System
sap_login("VA5", "700")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "sm59"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").expandNode ("          6")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "          1"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").selectItem ("         84","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").ensureVisibleHorizontalItem ("         84","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = ("         71")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").doubleClickItem ("         84","&Hierarchy")
session.findById("wnd[0]/tbar[1]/btn[27]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,139)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("vertex", "va5_700")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)


#DE5-200 System
sap_login("DE5", "200")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "sm59"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").expandNode ("          6")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "          1"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").selectItem ("         92","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").ensureVisibleHorizontalItem ("         92","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "         79"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").doubleClickItem ("         92","&Hierarchy")
session.findById("wnd[0]/tbar[1]/btn[27]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,131)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("vertex", "de5_200")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

#CV5-700 System
sap_login("CV5", "700")
session = create_session()
pyautogui.hotkey('win','up')

session.findById("wnd[0]/tbar[0]/okcd").text = "sm59"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").expandNode ("          5")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "          1"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").selectItem ("         67","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").ensureVisibleHorizontalItem ("         67","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "         54"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").doubleClickItem ("         67","&Hierarchy")
session.findById("wnd[0]/tbar[1]/btn[27]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,108)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("vertex", "cv5_700")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

#CD5-700 System
sap_login("CD5", "700")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "sm59"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").expandNode ("          6")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "          1"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").selectItem ("         63","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").ensureVisibleHorizontalItem ("         63","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "         50"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").doubleClickItem ("         63","&Hierarchy")
session.findById("wnd[0]/tbar[1]/btn[27]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,105)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("vertex", "cd5_700")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

#CP5-700 System
sap_login("CP5", "700")
session = create_session()
pyautogui.hotkey('win','up')
session.findById("wnd[0]").maximize()

session.findById("wnd[0]/tbar[0]/okcd").text = "sm59"
session.findById("wnd[0]").sendVKey (0)
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").expandNode ("          5")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "          1"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").selectItem ("         62","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").ensureVisibleHorizontalItem ("         62","&Hierarchy")
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").topNode = "         49"
session.findById("wnd[0]/usr/cntlSM59CNTL_AREA/shellcont/shell/shellcont[1]/shell[1]").doubleClickItem ("         62","&Hierarchy")
session.findById("wnd[0]/tbar[1]/btn[27]").press()
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (1,120)
session.findById("wnd[0]/usr/cntlGRID1/shellcont/shell").setRowSize (2,-1)
take_and_save_screenshot("vertex", "cp5_700")

session.findById("wnd[0]/tbar[0]/okcd").text = "/nex"
session.findById("wnd[0]").sendVKey (0)

base_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Screenshots'))

vertex_screenshot_names = [
    "cd5_700.png", "cp5_700.png", "cv5_700.png", "de5_200.png",
    "va5_700.png"
]

all_screenshot_paths = [
    f"{base_path}/vertex/{screenshot}" for screenshot in vertex_screenshot_names
]

today = date.today().strftime("%d.%m.%Y")
template_path = os.path.abspath(os.path.join(current_directory, 'Files','Vertex','Vertex_Monitoring_Template.docx'))
documents_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Documents'))
if not os.path.exists(documents_path):
    os.makedirs(documents_path)

output_path = os.path.abspath(os.path.join(current_directory, 'Scripts_run', 'Manual_Monitoring','Documents',f'Vertex_Monitoring_{today}.docx'))
doc = Document(template_path)

for i, screenshot_file in enumerate(all_screenshot_paths):
    placeholder = f"#IMAGE{i+1}#"
    for para in doc.paragraphs:
            if placeholder in para.text:
                index = para.text.find(placeholder)
                para.text = para.text.replace(placeholder,"")
                run = para.add_run()
                run.add_picture(screenshot_file, width=Inches(6))
                run.add_break()
                break

today = date.today().strftime("(%d/%m/%Y)")
for para in doc.paragraphs:
    if "#DATE#" in para.text:
        for run in para.runs:
            if "#DATE#" in run.text:
                run.text = run.text.replace("#DATE#", "")
                run.bold = True
                break
        run = para.add_run(today)
        run.bold = True


doc.save(output_path)

print(f"Word document with screenshots saved at {output_path}")